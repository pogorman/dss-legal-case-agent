"""
Sanitize source Word documents by replacing real case data with synthetic data.
Maps the real Erickson/Spartanburg case to the synthetic Webb/Holloway Case 1.

Usage: python docs/sanitize-docs.py
"""

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os
import glob

# =============================================================
# Replacement mapping: Real data → Synthetic data
# =============================================================
REPLACEMENTS = {
    # Case numbers
    "2025-DR-42-1286": "2024-DR-42-0892",
    "2023-DR-42-2003": "2024-DR-42-0893",
    "2023SPA00144": "2024SPA00892",

    # Parents
    "Kelle L. Erickson": "Dena Holloway",
    "Kelle Erickson": "Dena Holloway",
    "KELLE L. ERICKSON": "DENA HOLLOWAY",
    "KELLE ERICKSON": "DENA HOLLOWAY",
    "Adam Erickson": "Marcus Webb",
    "ADAM ERICKSON": "MARCUS WEBB",
    "Erickson": "Webb",
    "ERICKSON": "WEBB",

    # Child
    "Lydia Erickson": "Jaylen Webb",
    "LYDIA ERICKSON": "JAYLEN WEBB",
    "Lydia": "Jaylen",
    "LYDIA": "JAYLEN",
    "5/3/2019": "4/10/2021",
    "May 3, 2019": "April 10, 2021",

    # DSS Caseworker
    "Walela McDaniel": "Renee Dawson",
    "WALELA MCDANIEL": "RENEE DAWSON",
    "McDaniel": "Dawson",

    # Plaintiff attorney
    "Kathryn J Walsh": "Jennifer M. Torres",
    "Kathryn Walsh": "Jennifer Torres",
    "KATHRYN J WALSH": "JENNIFER M. TORRES",
    "KATHRYN WALSH": "JENNIFER TORRES",
    "454 South Anderson Road": "180 Magnolia Street",
    "York, SC 29745": "Spartanburg, SC 29306",

    # Defense attorneys
    "Tim Edwards, Esq.": "David Chen, Esq.",
    "Tim Edwards": "David Chen",
    "Shawn Campbell, Esq.": "Rachel Simmons, Esq.",
    "Shawn Campbell": "Rachel Simmons",

    # GAL
    "Jamia Foster, Esq.": "Karen Milford, Esq.",
    "Jamia Foster": "Karen Milford",
    "JAMIA FOSTER": "KAREN MILFORD",

    # Dates from real case (Feb 2026 filing dates → June 2024 synthetic dates)
    "February 17, 2026": "June 14, 2024",
    "2/17/2026": "6/14/2024",
}


def replace_in_paragraph(paragraph, replacements):
    """Replace text in a paragraph, handling runs carefully."""
    full_text = paragraph.text
    changed = False
    for old, new in replacements.items():
        if old in full_text:
            full_text = full_text.replace(old, new)
            changed = True
    if changed:
        # Preserve formatting of the first run, replace all text
        if paragraph.runs:
            # Store first run's formatting
            for i, run in enumerate(paragraph.runs):
                if i == 0:
                    run.text = full_text
                else:
                    run.text = ""
        else:
            paragraph.text = full_text
    return changed


def replace_in_table(table, replacements):
    """Replace text in all cells of a table."""
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if replace_in_paragraph(paragraph, replacements):
                    count += 1
    return count


def sanitize_docx(filepath, replacements):
    """Open a .docx file, replace all real data with synthetic data, save."""
    doc = Document(filepath)
    count = 0

    # Replace in body paragraphs
    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph, replacements):
            count += 1

    # Replace in tables
    for table in doc.tables:
        count += replace_in_table(table, replacements)

    # Replace in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.paragraphs:
                for paragraph in header.paragraphs:
                    if replace_in_paragraph(paragraph, replacements):
                        count += 1
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.paragraphs:
                for paragraph in footer.paragraphs:
                    if replace_in_paragraph(paragraph, replacements):
                        count += 1

    # Save back to same file
    doc.save(filepath)
    return count


def main():
    docs_dir = os.path.dirname(os.path.abspath(__file__))
    docx_files = glob.glob(os.path.join(docs_dir, "*.docx"))

    if not docx_files:
        print("No .docx files found in docs/ folder.")
        return

    print(f"Found {len(docx_files)} Word document(s) to sanitize.\n")

    # Sort replacements by length (longest first) to avoid partial matches
    sorted_replacements = dict(
        sorted(REPLACEMENTS.items(), key=lambda x: len(x[0]), reverse=True)
    )

    for filepath in docx_files:
        filename = os.path.basename(filepath)
        count = sanitize_docx(filepath, sorted_replacements)
        print(f"  {filename}: {count} paragraph(s) updated")

    # Rename files to use synthetic case number
    print("\nRenaming files to use synthetic case number...")
    for filepath in docx_files:
        filename = os.path.basename(filepath)
        new_filename = filename.replace("2023SPA00144", "2024SPA00892")
        if new_filename != filename:
            new_path = os.path.join(docs_dir, new_filename)
            os.rename(filepath, new_path)
            print(f"  {filename} -> {new_filename}")

    print("\nDone. All documents sanitized with synthetic data.")


if __name__ == "__main__":
    main()
