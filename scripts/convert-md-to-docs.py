"""
Convert all markdown files in sharepoint-docs/ to .docx and .pdf.
Outputs are saved alongside the source .md files.

Usage: python scripts/convert-md-to-docs.py
"""

import os
import re
import markdown
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from html.parser import HTMLParser


SHAREPOINT_DIR = os.path.join(os.path.dirname(__file__), "..", "sharepoint-docs")


class HTMLToText(HTMLParser):
    """Minimal HTML-to-text extractor for PDF generation."""

    def __init__(self):
        super().__init__()
        self.lines = []
        self.current = ""
        self.in_bold = False
        self.in_italic = False

    def handle_starttag(self, tag, attrs):
        if tag in ("h1", "h2", "h3", "h4"):
            self.flush()
            self.lines.append(("heading", tag, ""))
        elif tag == "p":
            self.flush()
        elif tag in ("strong", "b"):
            self.in_bold = True
        elif tag in ("em", "i"):
            self.in_italic = True
        elif tag == "br":
            self.flush()
        elif tag == "li":
            self.current += "  - "
        elif tag == "hr":
            self.flush()
            self.lines.append(("text", "---"))

    def handle_endtag(self, tag):
        if tag in ("h1", "h2", "h3", "h4"):
            self.lines.append(("heading", tag, self.current.strip()))
            self.current = ""
        elif tag == "p":
            self.flush()
        elif tag in ("strong", "b"):
            self.in_bold = False
        elif tag in ("em", "i"):
            self.in_italic = False
        elif tag == "li":
            self.flush()
        elif tag == "blockquote":
            self.flush()

    def handle_data(self, data):
        self.current += data

    def flush(self):
        text = self.current.strip()
        if text:
            self.lines.append(("text", text))
        self.current = ""

    def get_lines(self):
        self.flush()
        return self.lines


def md_to_docx(md_path, docx_path):
    """Convert a markdown file to .docx."""
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in md_text.split("\n"):
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("---"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("_" * 50)
            run.font.color.rgb = None
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("> "):
            p = doc.add_paragraph(stripped[2:])
            p.style = doc.styles["Quote"] if "Quote" in [s.name for s in doc.styles] else doc.styles["Normal"]
            p.paragraph_format.left_indent = Inches(0.5)
        elif stripped.startswith("|"):
            # Simple table row — just add as text
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells and all(c.replace("-", "").replace(":", "") == "" for c in cells):
                continue  # Skip separator rows
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run("  |  ".join(cells))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            # Handle inline bold/italic
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|_.*?_)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)

    doc.save(docx_path)


def md_to_pdf(md_path, pdf_path):
    """Convert a markdown file to .pdf."""
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    html = markdown.markdown(md_text, extensions=["tables"])

    parser = HTMLToText()
    parser.feed(html)
    lines = parser.get_lines()

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    for item in lines:
        if item[0] == "heading":
            tag = item[1]
            text = item[2]
            sizes = {"h1": 18, "h2": 15, "h3": 13, "h4": 11}
            size = sizes.get(tag, 11)
            pdf.set_font("Helvetica", "B", size)
            pdf.ln(4)
            pdf.multi_cell(0, size * 0.6, text.encode("latin-1", "replace").decode("latin-1"))
            pdf.ln(2)
        else:
            text = item[1]
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, text.encode("latin-1", "replace").decode("latin-1"))
            pdf.ln(1)

    pdf.output(pdf_path)


def main():
    count = 0
    for root, dirs, files in os.walk(SHAREPOINT_DIR):
        for fname in sorted(files):
            if not fname.endswith(".md"):
                continue

            md_path = os.path.join(root, fname)
            base = os.path.splitext(fname)[0]
            docx_path = os.path.join(root, base + ".docx")
            pdf_path = os.path.join(root, base + ".pdf")

            print(f"Converting {os.path.relpath(md_path, SHAREPOINT_DIR)}...")
            md_to_docx(md_path, docx_path)
            md_to_pdf(md_path, pdf_path)
            count += 1

    print(f"\nDone. Converted {count} files ({count} .docx + {count} .pdf).")


if __name__ == "__main__":
    main()
