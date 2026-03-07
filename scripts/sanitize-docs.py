"""
Sanitize source Word documents by replacing real case data with synthetic data.

Usage: python scripts/sanitize-docs.py

IMPORTANT: The REPLACEMENTS dict below must be populated with the real-to-synthetic
mapping before running. The mapping is not stored in version control to protect PII.
Copy the mapping from the local-only reference file before use.
"""

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os
import glob

# =============================================================
# Replacement mapping: Real data → Synthetic data
# NOT stored in version control — populate from local reference
# =============================================================
REPLACEMENTS = {
    # Populate this dict with real → synthetic mappings before running.
    # Example:
    # "Real Name": "Synthetic Name",
    # "Real Case Number": "Synthetic Case Number",
}


def replace_in_paragraph(paragraph, replacements):
    """Replace text in a paragraph, handling runs carefully."""
    full_text = paragraph.text
    changed = False
    for old, new in replacements.items():
        if old in full_text:
            full_text = full_text.replace(old, new)
            changed = True
    if changed:
        # Preserve formatting of the first run, replace all text
        if paragraph.runs:
            # Store first run's formatting
            for i, run in enumerate(paragraph.runs):
                if i == 0:
                    run.text = full_text
                else:
                    run.text = ""
        else:
            paragraph.text = full_text
    return changed


def replace_in_table(table, replacements):
    """Replace text in all cells of a table."""
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if replace_in_paragraph(paragraph, replacements):
                    count += 1
    return count


def sanitize_docx(filepath, replacements):
    """Open a .docx file, replace all real data with synthetic data, save."""
    doc = Document(filepath)
    count = 0

    # Replace in body paragraphs
    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph, replacements):
            count += 1

    # Replace in tables
    for table in doc.tables:
        count += replace_in_table(table, replacements)

    # Replace in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.paragraphs:
                for paragraph in header.paragraphs:
                    if replace_in_paragraph(paragraph, replacements):
                        count += 1
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.paragraphs:
                for paragraph in footer.paragraphs:
                    if replace_in_paragraph(paragraph, replacements):
                        count += 1

    # Save back to same file
    doc.save(filepath)
    return count


def main():
    if not REPLACEMENTS:
        print("ERROR: REPLACEMENTS dict is empty.")
        print("Populate it with the real-to-synthetic mapping from the local reference file.")
        return

    docs_dir = os.path.dirname(os.path.abspath(__file__))
    docx_files = glob.glob(os.path.join(docs_dir, "*.docx"))

    if not docx_files:
        print("No .docx files found in docs/ folder.")
        return

    print(f"Found {len(docx_files)} Word document(s) to sanitize.\n")

    # Sort replacements by length (longest first) to avoid partial matches
    sorted_replacements = dict(
        sorted(REPLACEMENTS.items(), key=lambda x: len(x[0]), reverse=True)
    )

    for filepath in docx_files:
        filename = os.path.basename(filepath)
        count = sanitize_docx(filepath, sorted_replacements)
        print(f"  {filename}: {count} paragraph(s) updated")

    print("\nDone. All documents sanitized with synthetic data.")


if __name__ == "__main__":
    main()
